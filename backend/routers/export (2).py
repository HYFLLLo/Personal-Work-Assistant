"""
报告导出路由 - 支持 PDF、Word、TXT、Markdown 格式导出
"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Literal
import io
from datetime import datetime

router = APIRouter()


class ExportRequest(BaseModel):
    content: str
    format: Literal["pdf", "word", "txt", "markdown"]
    title: str = "报告"


def export_to_txt(content: str) -> bytes:
    """导出为 TXT 格式"""
    return content.encode('utf-8')


def export_to_markdown(content: str) -> bytes:
    """导出为 Markdown 格式"""
    # 添加 Markdown 标题
    md_content = f"# 报告\n\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n---\n\n{content}"
    return md_content.encode('utf-8')


def export_to_word(content: str, title: str) -> bytes:
    """导出为 Word 格式"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # 设置默认中文字体
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 添加标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置标题字体
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 添加生成时间
        time_para = doc.add_paragraph()
        time_run = time_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        time_run.font.name = 'Microsoft YaHei'
        time_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        time_run.font.size = Pt(10)
        time_run.font.color.rgb = RGBColor(128, 128, 128)
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分隔线
        doc.add_paragraph("_" * 50)
        
        # 处理内容并添加到文档
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检测标题（以 # 开头）
            if line.startswith('# '):
                h = doc.add_heading(line.replace('# ', ''), level=1)
                for run in h.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            elif line.startswith('## '):
                h = doc.add_heading(line.replace('## ', ''), level=2)
                for run in h.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            elif line.startswith('### '):
                h = doc.add_heading(line.replace('### ', ''), level=3)
                for run in h.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            elif line.startswith('#### '):
                h = doc.add_heading(line.replace('#### ', ''), level=4)
                for run in h.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            # 检测列表项
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                p = doc.add_paragraph(line[3:], style='List Number')
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            else:
                # 普通段落
                para = doc.add_paragraph(line)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        raise HTTPException(status_code=500, detail="Word导出功能需要安装python-docx库")


def export_to_pdf(content: str, title: str) -> bytes:
    """导出为 PDF 格式"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        
        # 注册中文字体
        try:
            pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
            chinese_font = 'SimSun'
        except:
            try:
                pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
                chinese_font = 'SimSun'
            except:
                chinese_font = 'Helvetica'
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # 创建样式
        styles = getSampleStyleSheet()
        
        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=20,
            textColor='#1a1a1a',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        # 时间样式
        time_style = ParagraphStyle(
            'TimeStyle',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            textColor='#666666',
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        # 正文样式
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=11,
            leading=18,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )
        
        # 标题1样式
        h1_style = ParagraphStyle(
            'CustomH1',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=16,
            textColor='#2c3e50',
            spaceAfter=12,
            spaceBefore=12
        )
        
        # 标题2样式
        h2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontName=chinese_font,
            fontSize=14,
            textColor='#34495e',
            spaceAfter=10,
            spaceBefore=10
        )
        
        # 标题3样式
        h3_style = ParagraphStyle(
            'CustomH3',
            parent=styles['Heading3'],
            fontName=chinese_font,
            fontSize=12,
            textColor='#7f8c8d',
            spaceAfter=8,
            spaceBefore=8
        )
        
        # 构建文档内容
        story = []
        
        # 添加标题
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.1*inch))
        
        # 添加时间
        story.append(Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", time_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 处理内容
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1*inch))
                continue
            
            # 转义 HTML 特殊字符
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # 检测标题
            if line.startswith('# '):
                story.append(Paragraph(line.replace('# ', ''), h1_style))
            elif line.startswith('## '):
                story.append(Paragraph(line.replace('## ', ''), h2_style))
            elif line.startswith('### '):
                story.append(Paragraph(line.replace('### ', ''), h3_style))
            elif line.startswith('#### '):
                story.append(Paragraph(line.replace('#### ', ''), h3_style))
            # 检测列表项
            elif line.startswith('- ') or line.startswith('* '):
                story.append(Paragraph(f"• {line[2:]}", body_style))
            elif line[0:3] in ['1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ']:
                story.append(Paragraph(f"{line[0]}. {line[3:]}", body_style))
            else:
                story.append(Paragraph(line, body_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF导出功能需要安装reportlab库")


@router.post("/export")
async def export_report(request: ExportRequest):
    """
    导出报告为指定格式
    """
    try:
        if request.format == "txt":
            content = export_to_txt(request.content)
            media_type = "text/plain"
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        elif request.format == "markdown":
            content = export_to_markdown(request.content)
            media_type = "text/markdown"
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        
        elif request.format == "word":
            content = export_to_word(request.content, request.title)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        elif request.format == "pdf":
            content = export_to_pdf(request.content, request.title)
            media_type = "application/pdf"
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        else:
            raise HTTPException(status_code=400, detail="不支持的导出格式")
        
        return StreamingResponse(
            io.BytesIO(content),
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")
